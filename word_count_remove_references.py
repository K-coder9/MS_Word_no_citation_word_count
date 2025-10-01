from docx.api import Document
import re
def take_out_bracket(display_text_string):
    word_count = 0
    #find all the words inside of bracketss and put into a string array
    res = re.sub(r"\(.*?\)",'',display_text_string) 
    #take out the strings and upload the new paragraph as a new doc
    #print(res)
    split = res.split(' ')
    #get rid of punctuation 
    punc_counter = 0
    for i in range(len(split)):
        if split[i] == '.':
            punc_counter = punc_counter + 1
    #find how many full stops there are 
    for j in range(punc_counter):
        split.remove('.')
    #print(split)
    if split[0] == '': # this corrects so an empty paragraph does not contribute to word count/ new line
        word_count = 0
    else: 
     word_count = len(split)
    return word_count


#ask user for document name 
document_name = input('Please enter your document name: ')
name_with_extension = document_name + '.docx'
document = Document(str(name_with_extension))
#count the number of paragraphs in the word document
no_paragraphs = len(document.paragraphs)
#get text from current word document
word_count = 0
for j in range(no_paragraphs):
    display_text = document.paragraphs[j].text
#convert to string 
    display_text_string = str(display_text)
    word_count = take_out_bracket(display_text_string) + word_count
    
print(f'The word count is: {word_count}')
